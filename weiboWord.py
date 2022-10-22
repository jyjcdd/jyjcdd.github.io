# 代码运行环境：
# Python 3.x
# 代码运行所需库文件：
# selenium + docx

from selenium import webdriver
import time
import os
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from selenium.webdriver.common.by import By


class Weibo_spider:
    # Excel文件名
    wordName = {
        'jydd': '就业大队',
        'xqxz': '消歧小组',
        'hwzs': '海外之声'
    }

    def __init__(self, word_type, year, month, page, username, password):
        page = int(page)
        year = str(year)
        if len(month) == 1:
            month = '0' + str(month)

        # 账户信息
        self.username = username
        self.password = password
        # 页面信息
        #self.homeUrl = 'https://weibo.com/u/5327831786/home'
        #self.baseUrl = f'https://weibo.com/5327831786/profile?is_all=1&stat_date={year}{month}'
        #煎茶小队
        self.homeUrl = 'https://weibo.com/u/7722146918/home'
        self.baseUrl = f'https://weibo.com/7722146918/profile?is_all=1&stat_date={year}{month}'
        self.totalPageNum = page
        # 内容关键字
        if word_type == 'hwzs':
            self.keyWord = f'海外之声 {year}.{month}'

        # 爬虫
        self.driver = webdriver.Chrome()
        self.cur_data = []
        self.data_list = []
        # Excel表信息
        self.wordType = word_type
        self.wordName = f'{Weibo_spider.wordName[word_type]}微博{year}年{month}月数据'
        self.file = None
        self.initWord()

    def initWord(self):
        try:
            # 当前目录下是否有num.txt和相应word文件
            self.file = Document(f'{self.wordName}.docx')
            self.operate_file(f'{self.wordName}num.txt', 'r')
            print('txt & excel文件都存在')
        except:
            print('创建文件')
            self.file = Document()
            # 创建excel文件
            self.file.save(f'{self.wordName}.docx')
            # 创建num.txt并写入当前页码
            self.operate_file(f'{self.wordName}num.txt', 'w')

        # 设置字体
        self.file.styles['Normal'].font.name = u'宋体'
        self.file.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'宋体')

    def operate_file(self, filename, action):
        with open(filename, action) as f:
            # 写入页码记录
            if action == 'w':
                f.write(str(self.totalPageNum))
            elif action == 'r':
                # 将要抓的页码赋值
                self.totalPageNum = int(f.read())

    # 登录微博
    def login(self):
        # 访问微博首页
        self.driver.implicitly_wait(10)
        self.driver.get('https://weibo.com')
        time.sleep(5)
        try:#旧版页面可以做这个尝试
            # 登录账号
            user_input = self.driver.find_element(By.id,'loginname')
            user_input.send_keys(self.username)
            psd_input = self.driver.find_element(
                By.CSS_SELECTOR,'.password .W_input')
            psd_input.send_keys(self.password)
            # 点击登录按钮
            submit_btn = self.driver.find_element(
                By.CSS_SELECTOR,'.W_btn_a.btn_32px')
            time.sleep(1)
            submit_btn.click()
            try:
                vail_input = self.driver.find_element(
                    By.CSS_SELECTOR,'[node-type="verifycode"]')
                vail_input.send_keys('')
                print('请输入验证码...')
                time.sleep(3)
                while True:
                    if self.driver.current_url != self.homeUrl:
                        time.sleep(1)
                    else:
                        break
            except:
                # 等待网页加载
                time.sleep(3)
        except:#新版就要等我扫码登陆
            time.sleep(15)


    # 将滚动条移动到页面的底部
    def scroll_page(self):
        # 执行js代码
        scroll_js = "window.scrollTo(0, document.body.scrollHeight + 100);"
        scroll_height_js = "return document.body.scrollHeight;"

        # 微博pc端每页需要两次滚动加载完所有当页数据
        while True:
            # 获取滚动之前的高度
            prev_height = self.driver.execute_script(scroll_height_js)
            # 将滚动条置底
            self.driver.execute_script(scroll_js)
            time.sleep(3)
            # 获取滚动之后的高度
            next_height = self.driver.execute_script(scroll_height_js)
            # try:
            #     # 比较滚动前后高度 若已经到底最退出循环
            #     pre_text = self.driver.find_element(
            #         By.CSS_SELECTOR,'.page.prev').text
            #     if pre_text == '上一页' or pre_text == '上一頁':
            #         break
            # except:
            #     try:
            #         next_text = self.driver.find_element(
            #             By.CSS_SELECTOR,'.page.next').text
            #         if next_text == '下一页' or next_text == '下一頁':
            #             break
            #     except:
            #         pass
            try:
                # 比较滚动前后高度 若已经到底最退出循环
                #js="return document.querySelectorAll('.more_txt')"
                #wd为Webdirver
                #more_text_list=self.driver.execute_script(js)
                more_text_list = self.driver.find_elements(
                    By.CSS_SELECTOR,'.more_txt')
                print(more_text_list)
                more_flag = False
                for more_text in more_text_list:
                    print(more_text.text)
                    if more_text.text == '查看更早微博a':
                        more_flag = True
                        break
                if more_flag:
                    break
            except:
                print("fail")
                pass

        # 等待滚动完成
        time.sleep(1)
        # 回到顶部
        self.driver.execute_script('window.scrollTo(0, 0);')
        time.sleep(1)

    # 访问当前页码的网页
    def visit_current_page(self, page):
        self.driver.implicitly_wait(10)
        current_url = self.baseUrl + '&page=' + str(page)
        self.driver.get(current_url)
        time.sleep(3)
        for i in range(2):
            if self.driver.current_url == current_url:
                break

            self.driver.get(current_url)
            time.sleep(3)

    # 展开全文
    def __spider_full_text(self, ele):
        # 是否有展开全文
        has_full_text = False
        has_full_text_num = 3
        while has_full_text_num > 0 and not has_full_text:
            try:
                full_text_ele = ele.find_element(
                    By.CSS_SELECTOR,'.WB_text_opt[action-type="fl_unfold"]')
                has_full_text = True
                # 展开全文
                full_text_click_num = 5
                full_text_need_click = True
                while full_text_click_num > 0 and full_text_need_click:
                    try:
                        full_text_ele.click()
                        print('展开全文!')
                        full_text_need_click = False

                    except Exception:
                        full_text_click_num -= 1
                        # 曝光，以使得当前链接可点击
                        self.driver.execute_script(
                            'window.scrollTo(0, window.scrollY + 50);')
                        time.sleep(1)

            except:
                has_full_text_num -= 1
                # 曝光，以使得当前链接可点击
                self.driver.execute_script(
                    'window.scrollTo(0, window.scrollY + 50);')
                time.sleep(1)

            if not has_full_text:
                print('无展开全文')

        time.sleep(3)

    # 抓取文案内容
    def __spider_content(self, ele):
        # 展开全文
        text_ele = ele.find_element(By.CSS_SELECTOR,'.WB_text')
        self.__spider_full_text(text_ele)
        try:
            # 展开全文
            full_text_ele = ele.find_element(
                By.CSS_SELECTOR,'[node-type="feed_list_content_full"]')
            full_text = full_text_ele.text[:-5]  # 去除收起全文
        except:
            # 无展开全文
            full_text_ele = ele.find_element(
                By.CSS_SELECTOR,'[node-type="feed_list_content"]')
            full_text = full_text_ele.text

        if self.wordType == 'hwzs':
            # 抓取图片信息
            try:
                pics_ele = ele.find_elements(
                    By.CSS_SELECTOR,
                    '.WB_detail>.WB_media_wrap .WB_media_a li')
                for i in range(0, len(pics_ele) - 1):
                    img_ele = pics_ele[i].find_element(
                        By.CSS_SELECTOR,
                        'img')
                    full_text = full_text + '\n' + img_ele.get_attribute('src')
                    print(f'full_text:{full_text}')

            except Exception:
                pass

            self.cur_data.append(full_text)

    # 将数据写入Excel表
    def __write_word(self):
        # 写入微博数据
        # 微博倒叙排序，此处再倒叙一次
        index = len(self.data_list) - 1
        while index >= 0:
            self.file.add_paragraph(self.data_list[index])
            index -= 1

        self.file.save(f'{self.wordName}.docx')
        self.operate_file(f'{self.wordName}num.txt', 'w')
        # 清空当页数据列表
        self.data_list = []

    # 抓取数据
    def spider_data(self):
        # 登录
        self.login()
        time.sleep(5)
        while self.totalPageNum >= 1:
            print('page:', self.totalPageNum)
            # 访问网页
            self.visit_current_page(self.totalPageNum)
            # 保证循环继续到下一页
            self.totalPageNum -= 1
            # 滚动
            self.scroll_page()
            # 抓取数据
            elements = self.driver.find_elements(
                By.CSS_SELECTOR,
                '.WB_cardwrap.WB_feed_type')
            for ele in elements:
                detail_ele = ele.find_element(
                    By.CSS_SELECTOR,
                    '.WB_feed_detail')
                text_ele = detail_ele.find_element(By.CSS_SELECTOR,'.WB_text')
                if self.keyWord in text_ele.text:
                    self.cur_data = []
                    # 抓取文案内容
                    self.__spider_content(detail_ele)
                    # 写入数据列表
                    self.data_list.append(self.cur_data)
                    print(f'cur_data:{self.cur_data}')

            print(f'data_list:{self.data_list}')
            # 写入本页数据
            self.__write_word()

        # 关闭driver
        self.driver.close()
        # 并删除txt文件
        try:
            os.remove(f'{self.wordName}num.txt')
            print('已删除num.txt')
        except:
            print('无num.txt')


# 第一个参数是要抓取数据的对应缩写
# 第二、三个参数是抓取的年份、月份
# 第三个参数是一共的页数（到微博查看）
spider_type = sys.argv[1]
year = sys.argv[2]
month = sys.argv[3]
page = sys.argv[4]
username = sys.argv[5]
password = sys.argv[6]
ws = Weibo_spider(spider_type, year, month, page, username, password)
ws.spider_data()
print('Done')
